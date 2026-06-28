import os
import pdfplumber
from docx import Document
from config import UPLOAD_FOLDER


class ResumeParser:
    """
    Parse resumes from PDF and DOCX formats.
    Extracts plain text for NLP processing.
    """

    def __init__(self):
        """Initialize the parser."""
        self.upload_folder = UPLOAD_FOLDER

    def parse_file(self, filename):
        """
        Parse a resume file and extract its text.

        Args:
            filename: Name of the file in the uploads folder

        Returns:
            Extracted text as a string, or None if parsing fails
        """
        file_path = os.path.join(self.upload_folder, filename)

        # Check if file exists
        if not os.path.exists(file_path):
            print(f"File not found: {file_path}")
            return None

        # Get file extension
        _, ext = os.path.splitext(filename)
        ext = ext.lower()

        # Parse based on file type
        if ext == '.pdf':
            return self.parse_pdf(file_path)
        elif ext == '.docx':
            return self.parse_docx(file_path)
        else:
            print(f"Unsupported file format: {ext}")
            return None

    def parse_pdf(self, file_path):
        """
        Extract text from a PDF file using pdfplumber.

        HOW IT WORKS:
        1. Opens the PDF file
        2. Iterates through each page
        3. Extracts text from each page
        4. Combines all text into a single string

        Args:
            file_path: Full path to the PDF file

        Returns:
            Extracted text as a string
        """
        text = ""

        try:
            with pdfplumber.open(file_path) as pdf:
                # Process each page
                for page in pdf.pages:
                    # Extract text from page
                    page_text = page.extract_text()

                    if page_text:
                        # Clean up whitespace
                        page_text = page_text.strip()
                        text += page_text + "\n"

            return text.strip() if text.strip() else None

        except Exception as e:
            print(f"Error parsing PDF: {str(e)}")
            return None

    def parse_docx(self, file_path):
        """
        Extract text from a DOCX file using python-docx.

        HOW IT WORKS:
        1. Opens the DOCX file
        2. Iterates through paragraphs
        3. Extracts text from each paragraph
        4. Combines all text into a single string

        Args:
            file_path: Full path to the DOCX file

        Returns:
            Extracted text as a string
        """
        text = ""

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text += para_text + "\n"

            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        text += row_text + "\n"

            return text.strip() if text.strip() else None

        except Exception as e:
            print(f"Error parsing DOCX: {str(e)}")
            return None

    def get_file_info(self, filename):
        """
        Get basic information about the uploaded file.

        Args:
            filename: Name of the file

        Returns:
            Dictionary with file information
        """
        file_path = os.path.join(self.upload_folder, filename)

        if os.path.exists(file_path):
            stat = os.stat(file_path)
            return {
                'filename': filename,
                'size_bytes': stat.st_size,
                'size_kb': round(stat.st_size / 1024, 2),
                'exists': True
            }
        return {'exists': False}


# Utility function for direct use
def extract_text(filename):
    """
    Convenience function to extract text from a file.

    Args:
        filename: Name of the file in uploads folder

    Returns:
        Extracted text or None
    """
    parser = ResumeParser()
    return parser.parse_file(filename)


# Example usage for testing
if __name__ == '__main__':
    parser = ResumeParser()

    # Test with a sample file
    test_text = """
    John Doe
    Software Developer
    Email: john.doe@email.com
    Phone: +91-9876543210

    Skills: Python, JavaScript, React, Node.js, SQL, MongoDB

    Experience:
    - Senior Developer at ABC Corp (2020-Present)
    - Junior Developer at XYZ Inc (2018-2020)

    Education:
    - B.Tech in Computer Science, IIT Delhi (2014-2018)
    """

    print("Parser module loaded successfully!")
    print("\nSupported formats: PDF, DOCX")
